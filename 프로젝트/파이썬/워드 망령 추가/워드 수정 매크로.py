import asyncio
import openpyxl
from openpyxl.utils import get_column_letter
import os
from docx import Document
from docx.shared import RGBColor
import win32com.client as win32


result = 0
현재_폴더 = os.path.dirname(os.path.abspath(__file__))

async def 엑셀(반):
    global result
    excel_file = os.path.join(현재_폴더, '가라추가.xlsx') # 같은 디렉토리에 가라추가라는 xlsx 파일에 정리해야 함.
    wb = openpyxl.load_workbook(excel_file) # 엑셀 파일 열기
    sheet = wb.active
    column_name = get_column_letter(반 + 1)  # 열 이름을 가져옴
    column_values = []
    for cell in sheet[column_name][1:]:  # 첫 번째 행을 제외하고 순회 (반 이름 쓰는 곳)
        value = cell.value
        if value:  # 셀이 비어있지 않으면 [출]을 추가
            column_values.append(value + "[출]")

    result = ", " + ", ".join(column_values) # ', 이름[출]'을 result에 저장

    wb.close()


async def 파일_탐색_및_수정():
    global result, 현재_폴더
    for 반 in range(1, 51): # 1~50 폴더를 순회하게 함.
        폴더_경로 = os.path.join(현재_폴더, str(반))  # 폴더 경로는 현재 폴더에 지금 해당하는 반을 지정

        if os.path.exists(폴더_경로):          # 폴더가 존재하는지 확인 후
            await 엑셀(반)                     # 해당하는 반의 가라 학생 검색
            await 워드_수정(폴더_경로, result)  # result 값을 워드에 추가 하고 [출]으로 수정


async def 워드_수정(폴더_경로, result):
    for 파일명 in os.listdir(폴더_경로):
        if 파일명.endswith(".docx"):  # .docx 파일인 경우에만 처리
            워드_파일_경로 = os.path.join(폴더_경로, 파일명)

            doc = Document(워드_파일_경로)

           
            for table in doc.tables:  # 워드 파일 내 "출결현황" 칸 오른쪽에 result 값을 추가
                for row in table.rows:
                    for i, cell in enumerate(row.cells):
                        if cell.text == "출결현황" and i + 1 < len(row.cells):  # 출결현황이라는 텍스트 검색
                            next_cell = row.cells[i + 1]                       #  그 다음 셀 선택
                            next_cell.text = next_cell.text + result           # 맨 뒤에 result 추가
                        if "[결]" in cell.text:                                # [결]이라는 텍스트를 발견하면
                            cell.text = cell.text.replace("[결]", "[출]")      # [결]을 지우고 [출]로 변경
                        if "[-]" in cell.text:                                 # [-]이라는 텍스트를 발견하면
                            cell.text = cell.text.replace("[-]", "[출]")       # [-]를 지우고 [출]로 변경
                        if "[출]" in cell.text:                                # [출]이라는 텍스트를 발견하면
                            for paragraph in cell.paragraphs:
                                original_text = paragraph.text                 # 원래 단락의 텍스트를 저장
                                paragraph.clear()                              # 원래 단락의 텍스트를 지움
                                index = 0 
                                while index < len(original_text):              # "[출]" 문자열을 찾아 파란색으로 바꿈
                                    start_index = original_text.find("[출]", index)
                                    if start_index != -1:
                                        paragraph.add_run(original_text[index:start_index])  # "[출]" 문자열 전의 텍스트를 추가
                                        run = paragraph.add_run("[출]")                      # 파란색 "[출]" 문자열을 추가
                                        run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
                                        index = start_index + len("[출]")
                                    else:
                                        paragraph.add_run(original_text[index:])             # 남은 텍스트를 추가
                                        break

            doc.save(워드_파일_경로)

def 폴더_생성():
    global  현재_폴더
    for 반 in range(1, 51):
        폴더_경로 = os.path.join(현재_폴더, str(반))
        if not os.path.exists(폴더_경로):
            os.mkdir(폴더_경로) # 반 폴더가 없으면 폴더를 만듦



async def 변환_및_파일_수정():
    global 현재_폴더, 폴더_경로
    
    폴더_생성() # 먼저 폴더 생성

    for 반 in range(1, 51):
        폴더_경로 = os.path.join(현재_폴더, str(반))
        if os.path.exists(폴더_경로):
            await 엑셀(반)
            for 파일명 in os.listdir(폴더_경로): # .doc 파일을 .docx로 변환
                if 파일명.endswith(".doc"):
                    원본_파일_경로 = os.path.join(폴더_경로, 파일명)
                    변환된_파일_경로 = os.path.join(폴더_경로, 파일명 + "x")
                    word = win32.gencache.EnsureDispatch("Word.Application")
                    doc = word.Documents.Open(원본_파일_경로)
                    doc.SaveAs(변환된_파일_경로, FileFormat=16)
                    doc.Close()
                    os.remove(원본_파일_경로)   # 원본 .doc 파일 삭제
            await 워드_수정(폴더_경로, result)  # result 값을 전달
        print(f"{반}반 수정 완료")

async def 동작():
    await 변환_및_파일_수정()

asyncio.run(동작())