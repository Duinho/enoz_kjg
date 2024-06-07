import asyncio
import re
import sys
import openpyxl
from openpyxl.utils import get_column_letter
from playwright.async_api import async_playwright
import os
from docx import Document
from docx.shared import RGBColor
import win32com.client as win32

날짜 = '202405'
result = 0
폴더경로 = os.path.dirname(os.path.abspath(__file__))                               # 폴더경로를 코드가 있는 디렉토리로 저장 

async def 로그인():
    global page, new_handle     
    await page.goto('https://enozsw.enoz.kr/admin/')                           # 로그인 사이트로 이동
    await page.wait_for_load_state()                                                # 화면 바뀔 때까지 대기
    await asyncio.sleep(1)                                                          # 1초 대기
    await page.fill('input[name="tbAdminId"]', 'admin')                             # 아이디 및 비밀번호 입력하고 로그인
    await page.fill('input[name="tbAdminPass"]', 'admin55&&')
    await asyncio.sleep(1)
    await page.press('input[name="tbAdminPass"]', 'Enter')
    await asyncio.sleep(1)
    await page.goto('https://enozsw.enoz.kr/Admin/Class/ScheduleList.asp?ddlKeyField=a.group_no&tbKeyWord')    # 수강 관리로 이동
    await asyncio.sleep(1)
    await page.select_option('select[name="ddlTargetDate"]', value = 날짜)          # 날짜를 해당 기수로 변경
    await page.select_option('select[name="ddlKeyField"]', value ='Teacher')        # 아이디로 검색으로 바꾸기
    await asyncio.sleep(1)


async def 수정():
    global browser, page, new_handle,폴더경로  # 전역 변수로 사용
    links = await page.query_selector_all('a[href*="javascript:popStudyResult"]')
    for button in links: 
        await button.click()
        new_handle = None                                                                                        # 새로 열린 창 핸들 얻기
        while not new_handle:
            for handle in browser.contexts[0].pages:
                if handle != page:
                    new_handle = handle
                    break
                await asyncio.sleep(1)   
        await new_handle.wait_for_selector('a[href*="javascript:funcSave"]', timeout=30000)
        S_lists = await new_handle.query_selector_all('a[href="#on_list"]')
        print(S_lists)
        for S_list in S_lists:
            await S_list.click()    
            await asyncio.sleep(1)   
            await new_handle.select_option('select[name="study_chul"]', 'O')  
            await asyncio.sleep(1) 
    
    page.click('a[href*="javascript:funcSave"]')


async def 보고서다운():
    global browser, page, new_handle,폴더경로  # 전역 변수로 사용
    await 로그인()
    await 폴더_생성()
    for i in range(1, 11):                                                                                              # 50번 반복
        if i < 6:                                                                                                      # 26보다 작으면
            반 = 'GS_MW_'+ str(i).zfill(2)                                                                          # 반에 월수 몇 반이라고 저장
        else :                                                                                                          # 26 이상이면
            반 = 'GS_TT_' + str(i-5).zfill(2)                                                                      # 반에 화목 몇 반이라고 저장
                  
        폴더_경로 = os.path.join(폴더경로, str(i))                                                                       # 폴더_경로를 폴더경로의 숫자로 지정
        await page.fill('input[name="tbKeyWord"].font_blue', 반)                                                        # 입력창에 반을 입력
        await page.press('input[name="tbKeyWord"].font_blue', 'Enter')                                                  # 엔터로 검색
        await page.wait_for_selector('//a[contains(@class, "button_gray_small") and text()="보고서"]', timeout=30000)   # 회색 버튼의 보고서가 나올 때까지 대기
        await asyncio.sleep(2)                                                                                          # 2초 더 대기
        await 수정()
        buttons = await page.query_selector_all('//a[contains(@class, "button_gray_small") and text()="보고서"]')       # 회색 버튼의 보고서를 전부 찾음
        for button in buttons:                                                                                          # 찾은 개수만큼 반복
            await button.click()                                                                                        # 순차적으로 클릭
            await asyncio.sleep(1)                                                                                     
            new_handle = None                                                                                           # 새로 열린 창 핸들 얻기
            while not new_handle:
                for handle in browser.contexts[0].pages:
                    if handle != page:
                        new_handle = handle
                        break
            element = await new_handle.query_selector('xpath=//td[text()="담당강사명"]/following-sibling::td[1]')       
            if element:
                강사이름 = await new_handle.evaluate('element => element.textContent', element)                          # 강사이름에 담당강사명 저장
            element = await new_handle.query_selector('td.p_left5 .font_red b')
            if element:
                수업일시 = await new_handle.evaluate('element => element.textContent', element)                          # 수업일시에 수업날짜 저장
            
            file_name = f"{강사이름}_{수업일시}_보고서.doc"                                                               # 파일 이름 저장
            
            await new_handle.wait_for_selector('a.button_yellow', timeout=30000)                                        # 다운로드 버튼이 나올 때까지 대기
            dl_element = await new_handle.query_selector('a.button_yellow')                                             # 다운로드 버튼을 찾기
            if dl_element:
                async with new_handle.expect_download() as download_info:
                    await dl_element.click()  
                    download = await download_info.value
                await download.save_as(os.path.join(폴더_경로, file_name))                                              # 지정한 폴더에 다운로드하고 이름 바꾸기
                await new_handle.close()                                                                                # 새로 열린창 닫기
        
        red_buttons = await page.query_selector_all('//a[contains(@class, "button_red_small") and text()="보고서"]')    # 만약 빨간색으로 된 보고서 버튼이 있다면(미작성)
        for button in red_buttons:
            href_value = await button.get_attribute('href')
            match = re.search(r"popTeacherReport\('.*?', '(.*?)', '.*?'\)", href_value)                        
            if match:
                date_value = match.group(1)
                print(f"{강사이름}선생님은 {date_value}일 보고서를 작성하지 않으셨습니다.")                                # 강사이름과 미작성된 날짜 출력
    print('보고서 다운 완료')
    await browser.close()
                
async def 폴더_생성():
    global  폴더경로
    for 반 in range(1, 11):
        폴더_경로 = os.path.join(폴더경로, str(반))
        if not os.path.exists(폴더_경로):
            os.mkdir(폴더_경로) # 반 폴더가 없으면 폴더를 만듦   

async def 엑셀(반):
    global result,result1
    excel_file = os.path.join(폴더경로, '가라추가.xlsx') # 같은 디렉토리에 가라추가라는 xlsx 파일에 정리해야 함.
    wb = openpyxl.load_workbook(excel_file) # 엑셀 파일 열기
    sheet = wb.active
    column_name = get_column_letter(반)  # 열 이름을 가져옴
    column_values = []
    result = ""
    for cell in sheet[column_name][1:]:  # 첫 번째 행을 제외하고 순회 (반 이름 쓰는 곳)
        value = cell.value
        if value:  # 셀이 비어있지 않으면 [출]을 추가
            column_values.append(value + "[출]")
            result = ", " + ", ".join(column_values) # ', 이름[출]'을 result에 저장
    wb.close()
    excel_file = os.path.join(폴더경로, '취소자.xlsx') # 같은 디렉토리에 취소자라는 xlsx 파일에 정리해야 함.
    wb = openpyxl.load_workbook(excel_file) # 엑셀 파일 열기
    sheet = wb.active
    column_name = get_column_letter(반)  # 열 이름을 가져옴
    column_values = []
    for cell in sheet[column_name][1:]:  # 첫 번째 행을 제외하고 순회 (반 이름 쓰는 곳)
        value = cell.value
        if value:  # 셀이 비어있지 않으면 값을 리스트에 추가
            column_values.append(value)
    result1 = column_values  # 결과를 result1 변수에 저장
    wb.close()

async def 워드_수정(폴더_경로):
    global result,result1
    next_cell = None
    for 파일명 in os.listdir(폴더_경로):
        if 파일명.endswith(".docx"):  # .docx 파일인 경우에만 처리
            워드_파일_경로 = os.path.join(폴더_경로, 파일명)

            doc = Document(워드_파일_경로)

            for table in doc.tables:  
                for row in table.rows:
                    for i, cell in enumerate(row.cells):
                        if cell.text == "출결현황" and i + 1 < len(row.cells):  # 출결현황이라는 텍스트 검색
                            next_cell = row.cells[i + 1]                       #  그 다음 셀 선택
                            next_cell.text = next_cell.text + result           # 맨 뒤에 result 추가
                        if "[결]" in cell.text:                                # [결]이라는 텍스트를 발견하면
                            cell.text = cell.text.replace("[결]", "[출]")      # [결]을 지우고 [출]로 변경
                        if "[-]" in cell.text:                                 # [-]이라는 텍스트를 발견하면
                            cell.text = cell.text.replace("[-]", "[출]")       # [-]를 지우고 [출]로 변경
                    
                    # [출] 변경 작업이 끝난 후 [결] 추가 작업을 수행
                    if next_cell:
                        for name in result1:
                            if name not in next_cell.text:
                                next_cell.text += ", " + name + "[결]"
                
                # 색깔 변경 작업
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            original_text = paragraph.text  # 원래 단락의 텍스트를 저장
                            paragraph.clear()  # 원래 단락의 텍스트를 지움

                            index = 0
                            while index < len(original_text):
                                출_index = original_text.find("[출]", index)
                                결_index = original_text.find("[결]", index)

                                if 출_index == -1 and 결_index == -1:  # 더 이상 [출] 또는 [결] 문자열을 찾을 수 없는 경우
                                    paragraph.add_run(original_text[index:])
                                    break

                                if (출_index != -1 and 출_index < 결_index) or 결_index == -1:
                                    paragraph.add_run(original_text[index:출_index])
                                    run = paragraph.add_run("[출]")
                                    run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
                                    index = 출_index + len("[출]")
                                else:
                                    paragraph.add_run(original_text[index:결_index])
                                    run = paragraph.add_run("[결]")
                                    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                                    index = 결_index + len("[결]")
 
            doc.save(워드_파일_경로)

async def 변환_및_파일_수정():
    global 폴더경로, 폴더_경로
    for 반 in range(1, 51):
        폴더_경로 = os.path.join(폴더경로, str(반))
        if os.path.exists(폴더_경로):
            await 엑셀(반)
            for 파일명 in os.listdir(폴더_경로): # .doc 파일을 .docx로 변환
                if 파일명.endswith(".doc"):
                    원본_파일_경로 = os.path.join(폴더_경로, 파일명)
                    변환된_파일_경로 = os.path.join(폴더_경로, 파일명 + "x")
                    word = win32.gencache.EnsureDispatch("Word.Application")
                    doc = word.Documents.Open(원본_파일_경로)
                    doc.SaveAs(변환된_파일_경로, FileFormat=16)
                    doc.Close()
                    os.remove(원본_파일_경로)   # 원본 .doc 파일 삭제
            await 워드_수정(폴더_경로)  # result 값을 전달
        print(f'{반}반 수정 완료')
        
async def 동작():
    global browser, page, item, err  # 전역 변수로 사용
    async with async_playwright() as playwright: # 크로미움 브라우저 열기
        browser = await playwright.chromium.launch(
            headless=False,
        )
        page = await browser.new_page(accept_downloads=True) # 새로운 창이 열리면 page에 저장

        await 보고서다운()
        await 변환_및_파일_수정()

asyncio.run(동작())
