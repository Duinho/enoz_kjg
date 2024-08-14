import re
import os
import openpyxl
from openpyxl.utils import get_column_letter
from playwright.sync_api import sync_playwright
from docx import Document
from docx.shared import RGBColor
import win32com.client as win32
import time

result = 0
폴더경로 = os.path.dirname(os.path.abspath(__file__))  # 폴더경로를 코드가 있는 디렉토리로 저장 
excel_file_path = os.path.join(폴더경로, '보고서.xlsx')  # 엑셀 파일 경로

def 엑셀():
    global 로그인링크, 아이디, 비번, 수강관리링크, 강의날짜, 월수이름, 화목이름, 월수반수, 화목반수, 추가인원, 취소자

    wb = openpyxl.load_workbook(excel_file_path)
    ws = wb['정보']
    로그인링크 = ws.cell(row=1, column=17).value
    아이디 = ws.cell(row=2, column=17).value
    비번 = ws.cell(row=3, column=17).value
    수강관리링크 = ws.cell(row=4, column=17).value
    강의날짜 = str(ws.cell(row=5, column=17).value)  # 문자열로 저장
    월수이름 = ws.cell(row=6, column=17).value
    화목이름 = ws.cell(row=7, column=17).value
    월수반수 = int(ws.cell(row=8, column=17).value)  # 숫자로 저장
    화목반수 = int(ws.cell(row=9, column=17).value)  # 숫자로 저장
    추가인원 = wb['추가인원']
    취소자 = wb['취소자']
    wb.close()

def 엑셀값저장(반):
    global result, result1
    column_name = get_column_letter(반)  # 열 이름을 가져옴

    # 추가인원 시트에서 값 읽기
    column_values = []
    result = ""
    for cell in 추가인원[column_name][1:]:  # 첫 번째 행을 제외하고 순회 (반 이름 쓰는 곳)
        value = cell.value
        if value:  # 셀이 비어있지 않으면 [출]을 추가
            column_values.append(value + "[출]")
            result = ", " + ", ".join(column_values)  # ', 이름[출]'을 result에 저장

    # 취소자 시트에서 값 읽기
    column_values = []
    for cell in 취소자[column_name][1:]:  # 첫 번째 행을 제외하고 순회 (반 이름 쓰는 곳)
        value = cell.value
        if value:  # 셀이 비어있지 않으면 값을 리스트에 추가
            column_values.append(value)
    result1 = column_values  # 결과를 result1 변수에 저장

def 로그인(page):
    page.goto(로그인링크)  # 로그인 사이트로 이동
    page.wait_for_load_state()  # 화면 바뀔 때까지 대기
    page.fill('input[name="tbAdminId"]', 아이디)  # 아이디 및 비밀번호 입력하고 로그인
    page.fill('input[name="tbAdminPass"]', 비번)
    page.press('input[name="tbAdminPass"]', 'Enter')
    page.goto(수강관리링크)  # 수강 관리로 이동
    page.select_option('select[name="ddlTargetDate"]', value=강의날짜)  # 날짜를 해당 기수로 변경
    page.select_option('select[name="ddlKeyField"]', value='Teacher')  # 아이디로 검색으로 바꾸기
    time.sleep(1)

def 보고서다운(page, context):
    global 폴더경로  # 전역 변수로 사용
    로그인(page)
    for i in range(0, 월수반수 + 화목반수):  # 모든 반을 반복
        if i < 월수반수:
            반 = 월수이름 + str(i + 1).zfill(2)  # 반에 월수 몇 반이라고 저장
        else:
            반 = 화목이름 + str(i - 월수반수 + 1).zfill(2)  # 반에 화목 몇 반이라고 저장
        폴더_경로 = os.path.join(폴더경로, 반)
        if not os.path.exists(폴더_경로):
            os.mkdir(폴더_경로)  # 반 폴더가 없으면 폴더를 만듦

        page.fill('input[name="tbKeyWord"].font_blue', 반)  # 입력창에 반을 입력
        page.press('input[name="tbKeyWord"].font_blue', 'Enter')  # 엔터로 검색
        page.wait_for_selector('//a[contains(@class, "button_gray_small") and text()="보고서"]', timeout=30000)  # 회색 버튼의 보고서가 나올 때까지 대기
        time.sleep(3)
        buttons = page.query_selector_all('//a[contains(@class, "button_gray_small") and text()="보고서"]')  # 회색 버튼의 보고서를 전부 찾음
        for button in buttons:  # 찾은 개수만큼 반복
            with context.expect_page() as new_page_info:  # 새 페이지가 열리기를 대기
                button.click()
            new_page = new_page_info.value
            new_page.wait_for_load_state()
            element = new_page.query_selector('xpath=//td[text()="담당강사명"]/following-sibling::td[1]')
            if element:
                강사이름 = new_page.evaluate('element => element.textContent', element)  # 강사이름에 담당강사명 저장
            element = new_page.query_selector('td.p_left5 .font_red b')
            if element:
                수업일시 = new_page.evaluate('element => element.textContent', element)  # 수업일시에 수업날짜 저장

            file_name = f"{강사이름}_{수업일시}_보고서.doc"  # 파일 이름 저장

            new_page.wait_for_selector('a.button_yellow', timeout=30000)  # 다운로드 버튼이 나올 때까지 대기
            dl_element = new_page.query_selector('a.button_yellow')  # 다운로드 버튼을 찾기
            if dl_element:
                with new_page.expect_download() as download_info:
                    dl_element.click()
                    download = download_info.value
                download.save_as(os.path.join(폴더_경로, file_name))  # 지정한 폴더에 다운로드하고 이름 바꾸기
                new_page.close()  # 새로 열린창 닫기

        red_buttons = page.query_selector_all('//a[contains(@class, "button_red_small") and text()="보고서"]')  # 만약 빨간색으로 된 보고서 버튼이 있다면(미작성)
        for button in red_buttons:
            href_value = button.get_attribute('href')
            match = re.search(r"popTeacherReport\('.*?', '(.*?)', '.*?'\)", href_value)
            if match:
                date_value = match.group(1)
                print(f"{강사이름}선생님은 {date_value}일 보고서를 작성하지 않으셨습니다.")  # 강사이름과 미작성된 날짜 출력
    print('보고서 다운 완료')
    browser.close()

def 워드_수정(폴더_경로):
    global result, result1
    for 파일명 in os.listdir(폴더_경로):
        if 파일명.endswith(".docx"):  # .docx 파일인 경우에만 처리
            워드_파일_경로 = os.path.join(폴더_경로, 파일명)

            doc = Document(워드_파일_경로)

            for table in doc.tables:
                for row in table.rows:
                    for i, cell in enumerate(row.cells):
                        if cell.text == "출결현황" and i + 1 < len(row.cells):  # 출결현황이라는 텍스트 검색
                            next_cell = row.cells[i + 1]  # 그 다음 셀 선택
                            next_cell_text = next_cell.text

                            # 먼저 [결]과 [-]를 [출]로 변경
                            next_cell_text = next_cell_text.replace("[결]", "[출]")
                            next_cell_text = next_cell_text.replace("[-]", "[출]")

                            # 취소자 이름 뒤의 [출]을 [결]로 수정
                            for name in result1:
                                pattern = re.compile(rf"({re.escape(name)}\[출\])")
                                next_cell_text = pattern.sub(rf"{name}[결]", next_cell_text)

                            next_cell.text = next_cell_text + result  # 맨 뒤에 result 추가

                # 색깔 변경 작업
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            original_text = paragraph.text  # 원래 단락의 텍스트를 저장
                            paragraph.clear()  # 원래 단락의 텍스트를 지움

                            index = 0
                            while index < len(original_text):
                                출_index = original_text.find("[출]", index)
                                결_index = original_text.find("[결]", index)

                                if 출_index == -1 and 결_index == -1:  # 더 이상 [출] 또는 [결] 문자열을 찾을 수 없는 경우
                                    paragraph.add_run(original_text[index:])
                                    break

                                if (출_index != -1 and 출_index < 결_index) or 결_index == -1:
                                    paragraph.add_run(original_text[index:출_index])
                                    run = paragraph.add_run("[출]")
                                    run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
                                    index = 출_index + len("[출]")
                                else:
                                    paragraph.add_run(original_text[index:결_index])
                                    run = paragraph.add_run("[결]")
                                    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                                    index = 결_index + len("[결]")

            doc.save(워드_파일_경로)

def 변환_및_파일_수정():
    global 폴더경로, 폴더_경로
    for i in range(0, 월수반수 + 화목반수):
        if i < 월수반수:
            반 = 월수이름 + str(i + 1).zfill(2)  # 반에 월수 몇 반이라고 저장
        else:
            반 = 화목이름 + str(i - 월수반수 + 1).zfill(2)  # 반에 화목 몇 반이라고 저장
        폴더_경로 = os.path.join(폴더경로, 반)
        if os.path.exists(폴더_경로):
            엑셀값저장(i + 1)  # 1부터 시작하도록 조정
            for 파일명 in os.listdir(폴더_경로):  # .doc 파일을 .docx로 변환
                if 파일명.endswith(".doc"):
                    원본_파일_경로 = os.path.join(폴더_경로, 파일명)
                    변환된_파일_경로 = os.path.join(폴더_경로, 파일명 + "x")
                    word = win32.gencache.EnsureDispatch("Word.Application")
                    doc = word.Documents.Open(원본_파일_경로)
                    doc.SaveAs(변환된_파일_경로, FileFormat=16)
                    doc.Close()
                    os.remove(원본_파일_경로)  # 원본 .doc 파일 삭제
            워드_수정(폴더_경로)  # result 값을 전달
        print(f'{반}반 수정 완료')

def 워드_파일_병합(폴더경로, 저장_파일명):
    merged_document = Document()
    for 반 in range(0, 월수반수 + 화목반수):
        if 반 < 월수반수:
            반_이름 = 월수이름 + str(반 + 1).zfill(2)
        else:
            반_이름 = 화목이름 + str(반 - 월수반수 + 1).zfill(2)
        폴더_경로 = os.path.join(폴더경로, 반_이름)
        if os.path.exists(폴더_경로):
            for 파일명 in os.listdir(폴더_경로):
                if 파일명.endswith(".docx"):
                    파일_경로 = os.path.join(폴더_경로, 파일명)
                    sub_doc = Document(파일_경로)
                    for element in sub_doc.element.body:
                        merged_document.element.body.append(element)

    # 병합 후 '양식의 맨 아래' 문구와 빈 단락 제거
    for paragraph in merged_document.paragraphs:
        if '양식의 맨 아래' in paragraph.text:
            paragraph.text = paragraph.text.replace('양식의 맨 아래', '')

    # 빈 단락 제거
    merged_document.paragraphs[:] = [p for p in merged_document.paragraphs if p.text.strip()]

    merged_document.save(os.path.join(폴더경로, 저장_파일명))

def 동작():
    global browser, page  # 전역 변수로 사용
    with sync_playwright() as p:  # 크로미움 브라우저 열기
        browser = p.chromium.launch(headless=False)
        context = browser.new_context(accept_downloads=True)
        page = context.new_page()  # 새로운 창이 열리면 page에 저장
        엑셀()
        보고서다운(page, context)
        변환_및_파일_수정()
        워드_파일_병합(폴더경로, '합쳐진_보고서.docx')

동작()
